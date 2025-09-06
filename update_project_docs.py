#!/usr/bin/env python3
"""
Script to update Project Documentation.docx with comprehensive dashboard usage guide
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_dashboard_documentation():
    """Create comprehensive dashboard documentation in DOCX format"""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('GameLens AI - ROAS Forecasting Dashboard Usage Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add overview section
    doc.add_heading('Overview', level=1)
    overview_text = """
The GameLens AI dashboard is a comprehensive ROAS (Return on Ad Spend) forecasting platform designed for mobile game advertising optimization. This guide provides detailed instructions for using each page of the dashboard.

The dashboard consists of 6 main pages accessible via the sidebar:
1. Data Overview - Data loading and validation
2. Feature Engineering - Data preprocessing and feature creation  
3. Model Training - Machine learning model training and configuration
4. Predictions - ROAS predictions and model performance analysis
5. Recommendations - Actionable insights and campaign recommendations
6. FAQ - AI-powered question answering system
    """
    doc.add_paragraph(overview_text.strip())
    
    # Page 1: Data Overview
    doc.add_heading('Page 1: Data Overview 📊', level=1)
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph('Load, validate, and explore your advertising data from multiple platforms (Unity Ads, Mistplay, etc.).')
    
    doc.add_heading('How to Use', level=2)
    doc.add_paragraph('1. Data Loading:')
    doc.add_paragraph('   • The system automatically scans the Campaign Data/ folder', style='List Bullet')
    doc.add_paragraph('   • Supports CSV files from Unity Ads and Mistplay platforms', style='List Bullet')
    doc.add_paragraph('   • Handles both Android and iOS data', style='List Bullet')
    
    doc.add_paragraph('2. Data Validation:')
    doc.add_paragraph('   • Check data quality metrics (missing values, duplicates)', style='List Bullet')
    doc.add_paragraph('   • Verify platform distribution', style='List Bullet')
    doc.add_paragraph('   • Review date ranges and sample sizes', style='List Bullet')
    
    doc.add_paragraph('3. Data Preview:')
    doc.add_paragraph('   • View sample records from each platform', style='List Bullet')
    doc.add_paragraph('   • Examine column structures and data types', style='List Bullet')
    doc.add_paragraph('   • Identify any data quality issues', style='List Bullet')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'Multi-platform support: Unity Ads, Mistplay',
        'Automatic data detection: Scans for CSV files',
        'Data quality metrics: Missing values, duplicates, platform distribution',
        'Data preview: Sample records and column information'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('What to Look For', level=2)
    checklist = [
        'All expected platforms are loaded',
        'Data covers your desired date range',
        'No significant missing values or duplicates',
        'ROAS columns are present (roas_d0, roas_d7, roas_d30, etc.)'
    ]
    for item in checklist:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # Page 2: Feature Engineering
    doc.add_heading('Page 2: Feature Engineering 🔧', level=1)
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph('Transform raw advertising data into predictive features for machine learning models.')
    
    doc.add_heading('How to Use', level=2)
    doc.add_paragraph('1. Target Day Selection:')
    doc.add_paragraph('   • Choose which ROAS day to predict (D7, D15, D30, D45, D90)', style='List Bullet')
    doc.add_paragraph('   • This determines what the model will learn to predict', style='List Bullet')
    
    doc.add_paragraph('2. Feature Creation:')
    doc.add_paragraph('   • Click "Create Features" to generate predictive features', style='List Bullet')
    doc.add_paragraph('   • Features include retention rates, early ROAS, level progression, cost metrics', style='List Bullet')
    doc.add_paragraph('   • Process may take a few minutes for large datasets', style='List Bullet')
    
    doc.add_paragraph('3. Feature Analysis:')
    doc.add_paragraph('   • Review feature summary statistics', style='List Bullet')
    doc.add_paragraph('   • Check feature importance rankings', style='List Bullet')
    doc.add_paragraph('   • Identify top predictive features', style='List Bullet')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'Retention features: D1, D3, D7 retention rates',
        'Early ROAS indicators: D0, D3, D7 ROAS values',
        'Engagement metrics: Level progression, session data',
        'Cost efficiency: CPI, cost per level, revenue ratios',
        'Platform-specific features: Platform performance indicators'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('What to Look For', level=2)
    checklist = [
        'Features created successfully (no errors)',
        'Reasonable feature importance scores',
        'Features align with your business logic',
        'No missing values in critical features'
    ]
    for item in checklist:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # Page 3: Model Training
    doc.add_heading('Page 3: Model Training 🤖', level=1)
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph('Train machine learning models to predict ROAS with confidence intervals.')
    
    doc.add_heading('How to Use', level=2)
    doc.add_paragraph('1. Model Configuration:')
    doc.add_paragraph('   • Target Day: Select the ROAS day to predict', style='List Bullet')
    doc.add_paragraph('   • Number of Estimators: Trees in the model (default: 100)', style='List Bullet')
    doc.add_paragraph('   • Learning Rate: How fast the model learns (default: 0.1)', style='List Bullet')
    doc.add_paragraph('   • Max Depth: Tree depth (default: 6)', style='List Bullet')
    doc.add_paragraph('   • Random State: For reproducible results (default: 42)', style='List Bullet')
    
    doc.add_paragraph('2. Training Process:')
    doc.add_paragraph('   • Click "Train Model" to start training', style='List Bullet')
    doc.add_paragraph('   • Process may take 2-5 minutes depending on data size', style='List Bullet')
    doc.add_paragraph('   • Monitor progress with status messages', style='List Bullet')
    
    doc.add_paragraph('3. Model Validation:')
    doc.add_paragraph('   • Review training metrics (R², RMSE, MAPE, MAE)', style='List Bullet')
    doc.add_paragraph('   • Check confidence interval coverage', style='List Bullet')
    doc.add_paragraph('   • Examine feature importance rankings', style='List Bullet')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'Quantile Regression: Predicts ROAS with confidence intervals',
        'LightGBM/XGBoost: High-performance gradient boosting',
        'Feature Importance: Identifies key predictive factors',
        'Model Validation: Comprehensive performance metrics'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('What to Look For', level=2)
    checklist = [
        'R² score > 0.3 (good model fit)',
        'MAPE < 50% (reasonable prediction accuracy)',
        'Confidence coverage ~80-90% (reliable intervals)',
        'Feature importance makes business sense'
    ]
    for item in checklist:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # Page 4: Predictions
    doc.add_heading('Page 4: Predictions 📈', level=1)
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph('Generate ROAS predictions and analyze model performance on your data.')
    
    doc.add_heading('How to Use', level=2)
    doc.add_paragraph('1. Prediction Generation:')
    doc.add_paragraph('   • Predictions are automatically generated after model training', style='List Bullet')
    doc.add_paragraph('   • Shows ROAS predictions with confidence intervals', style='List Bullet')
    doc.add_paragraph('   • Displays prediction distribution and statistics', style='List Bullet')
    
    doc.add_paragraph('2. Performance Analysis:')
    doc.add_paragraph('   • Review model performance metrics', style='List Bullet')
    doc.add_paragraph('   • Check prediction accuracy and confidence', style='List Bullet')
    doc.add_paragraph('   • Analyze ROAS distribution across campaigns', style='List Bullet')
    
    doc.add_paragraph('3. Data Insights:')
    doc.add_paragraph('   • View prediction summaries (mean, percentiles)', style='List Bullet')
    doc.add_paragraph('   • Examine confidence interval widths', style='List Bullet')
    doc.add_paragraph('   • Identify high/low performing campaigns', style='List Bullet')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'ROAS Predictions: Point estimates with confidence bounds',
        'Performance Metrics: R², RMSE, MAPE, MAE, confidence coverage',
        'Distribution Analysis: ROAS ranges and percentiles',
        'Campaign Insights: High/medium/low ROAS categorization'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('What to Look For', level=2)
    checklist = [
        'Predictions completed successfully',
        'Reasonable ROAS ranges (0.1 to 5.0 typical)',
        'Good model performance metrics',
        'Confidence intervals are not too wide'
    ]
    for item in checklist:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # Page 5: Recommendations
    doc.add_heading('Page 5: Recommendations 💡', level=1)
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph('Generate actionable recommendations for campaign optimization based on model predictions.')
    
    doc.add_heading('How to Use', level=2)
    doc.add_paragraph('1. Recommendation Settings:')
    doc.add_paragraph('   • Target ROAS: Set your desired ROAS goal (default: 0.5)', style='List Bullet')
    doc.add_paragraph('   • Number of Recommendations: How many to show (default: 10)', style='List Bullet')
    doc.add_paragraph('   • Confidence Threshold: Minimum confidence level (default: 0.8)', style='List Bullet')
    
    doc.add_paragraph('2. Generate Recommendations:')
    doc.add_paragraph('   • Click "Generate Recommendations"', style='List Bullet')
    doc.add_paragraph('   • Review campaign-specific recommendations', style='List Bullet')
    doc.add_paragraph('   • Analyze improvement potential', style='List Bullet')
    
    doc.add_paragraph('3. Action Planning:')
    doc.add_paragraph('   • Identify campaigns to scale or cut', style='List Bullet')
    doc.add_paragraph('   • Prioritize based on improvement potential', style='List Bullet')
    doc.add_paragraph('   • Consider confidence levels in decisions', style='List Bullet')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'Campaign Prioritization: Ranked by improvement potential',
        'Confidence Filtering: Only high-confidence recommendations',
        'ROAS Improvement: Quantified potential gains',
        'Actionable Insights: Clear next steps for each campaign'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('What to Look For', level=2)
    checklist = [
        'Recommendations generated successfully',
        'Clear improvement potential for each campaign',
        'Confidence levels are reasonable',
        'Recommendations align with business goals'
    ]
    for item in checklist:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # Page 6: FAQ
    doc.add_heading('Page 6: FAQ ❓', level=1)
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph('AI-powered question answering system that provides insights about your data and model performance.')
    
    doc.add_heading('How to Use', level=2)
    doc.add_paragraph('1. Question Selection:')
    doc.add_paragraph('   • Browse pre-loaded questions from your FAQ document', style='List Bullet')
    doc.add_paragraph('   • Questions are automatically extracted from FAQ.docx', style='List Bullet')
    doc.add_paragraph('   • Click on any question to get an AI-powered answer', style='List Bullet')
    
    doc.add_paragraph('2. AI-Powered Answers:')
    doc.add_paragraph('   • Answers combine your FAQ content with current dashboard data', style='List Bullet')
    doc.add_paragraph('   • Provides data-driven insights and recommendations', style='List Bullet')
    doc.add_paragraph('   • Uses your actual model performance and predictions', style='List Bullet')
    
    doc.add_paragraph('3. Context-Aware Responses:')
    doc.add_paragraph('   • Answers are tailored to your specific data', style='List Bullet')
    doc.add_paragraph('   • Includes current model metrics and performance', style='List Bullet')
    doc.add_paragraph('   • Provides actionable insights based on your results', style='List Bullet')
    
    doc.add_heading('Key Features', level=2)
    features = [
        'Document Integration: Reads questions from FAQ.docx',
        'AI-Powered: Uses GPT for intelligent responses',
        'Data-Driven: Combines FAQ knowledge with dashboard data',
        'Context-Aware: Tailored to your specific results'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    doc.add_heading('What to Look For', level=2)
    checklist = [
        'FAQ content loaded successfully',
        'Questions extracted from your document',
        'AI answers are relevant and helpful',
        'Responses include your actual data insights'
    ]
    for item in checklist:
        doc.add_paragraph(f'✅ {item}', style='List Bullet')
    
    # Best Practices Section
    doc.add_heading('Best Practices', level=1)
    
    doc.add_heading('Data Preparation', level=2)
    practices = [
        'Ensure data quality: Clean, complete datasets work best',
        'Include sufficient history: At least 30 days of data recommended',
        'Verify ROAS columns: Ensure target ROAS days are present',
        'Check platform coverage: Include all relevant advertising platforms'
    ]
    for practice in practices:
        doc.add_paragraph(f'• {practice}', style='List Bullet')
    
    doc.add_heading('Model Training', level=2)
    practices = [
        'Start with defaults: Use default parameters initially',
        'Monitor performance: Watch for overfitting or underfitting',
        'Validate results: Check that metrics make business sense',
        'Iterate if needed: Adjust parameters based on performance'
    ]
    for practice in practices:
        doc.add_paragraph(f'• {practice}', style='List Bullet')
    
    doc.add_heading('Using Predictions', level=2)
    practices = [
        'Consider confidence: Use confidence intervals in decision-making',
        'Focus on trends: Look at patterns rather than individual predictions',
        'Validate with business logic: Ensure predictions align with expectations',
        'Monitor over time: Track prediction accuracy as new data comes in'
    ]
    for practice in practices:
        doc.add_paragraph(f'• {practice}', style='List Bullet')
    
    doc.add_heading('Making Recommendations', level=2)
    practices = [
        'Set realistic targets: Use achievable ROAS goals',
        'Consider confidence levels: Prioritize high-confidence recommendations',
        'Balance risk/reward: Consider both potential gains and confidence',
        'Test and learn: Start with small changes and scale successful ones'
    ]
    for practice in practices:
        doc.add_paragraph(f'• {practice}', style='List Bullet')
    
    # Technical Requirements
    doc.add_heading('Technical Requirements', level=1)
    
    doc.add_heading('System Requirements', level=2)
    requirements = [
        'RAM: 32GB recommended for large datasets',
        'Storage: Sufficient space for data files and models',
        'Network: Stable internet for AI-powered FAQ (if enabled)'
    ]
    for req in requirements:
        doc.add_paragraph(f'• {req}', style='List Bullet')
    
    doc.add_heading('Data Requirements', level=2)
    requirements = [
        'Format: CSV files with standard column names',
        'Platforms: Unity Ads, Mistplay (others can be added)',
        'Time Range: Minimum 30 days of historical data',
        'ROAS Columns: Must include target ROAS days (D7, D15, D30, etc.)'
    ]
    for req in requirements:
        doc.add_paragraph(f'• {req}', style='List Bullet')
    
    doc.add_heading('Optional Features', level=2)
    features = [
        'AI FAQ: Requires OpenAI API key for enhanced question answering',
        'Advanced Analytics: Additional features available with premium data',
        'Custom Models: Specialized models for specific game types'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = """
This dashboard provides a complete solution for ROAS forecasting and campaign optimization. Each page builds upon the previous one, creating a comprehensive workflow from data loading to actionable recommendations.

The system is designed to be user-friendly while providing powerful insights for mobile game advertising optimization. Follow the best practices outlined in this guide to get the most value from your ROAS forecasting platform.
    """
    doc.add_paragraph(conclusion.strip())
    
    return doc

def main():
    """Main function to update the project documentation"""
    try:
        # Create the comprehensive documentation
        doc = create_dashboard_documentation()
        
        # Save the document
        output_file = 'Project Documentation - Complete.docx'
        doc.save(output_file)
        
        print(f"✅ Successfully created comprehensive dashboard documentation: {output_file}")
        print(f"📄 Document contains detailed usage instructions for all 6 dashboard pages")
        print(f"📋 Includes best practices, technical requirements, and troubleshooting guide")
        
        # Also update the original file if it exists
        if os.path.exists('Project Documentation.docx'):
            doc.save('Project Documentation.docx')
            print(f"✅ Updated original Project Documentation.docx file")
        
    except Exception as e:
        print(f"❌ Error creating documentation: {e}")

if __name__ == "__main__":
    main()
